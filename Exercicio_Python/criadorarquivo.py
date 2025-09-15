from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Capa padrão Senac
doc.add_paragraph().add_run("SENAC\n").bold = True
capa = doc.paragraphs[0]
capa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph().add_run("Centro de Educação Profissional\nCurso: Administração\nDisciplina: Administração Geral\n").bold = False
for p in doc.paragraphs[1:]:
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Título
doc.add_paragraph().add_run("Taylorismo Digital: Reflexos do Pensamento Administrativo Clássico nas Organizações Contemporâneas").bold = True
titulo = doc.paragraphs[-1]
titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Adiciona o texto formatado
texto = '''
O pensamento administrativo clássico, representado pela Administração Científica de Frederick Taylor e pela Administração Geral de Henri Fayol, buscava organizar e racionalizar o trabalho. Taylor defendia a divisão de tarefas, estudo de tempos e movimentos e incentivos salariais, enquanto Fayol estruturava princípios de planejamento, organização, comando, coordenação e controle. Esses fundamentos foram essenciais para o avanço da produção industrial no início do século XX.

Nas organizações altamente tecnológicas contemporâneas, observa-se tanto a permanência quanto a adaptação desses princípios. A padronização e a mensuração propostas por Taylor estão presentes em algoritmos, métricas de desempenho e indicadores em tempo real. Já os princípios de Fayol se refletem em modelos de governança corporativa, compliance e planejamento estratégico. No entanto, diferentemente do passado, essas práticas convivem com demandas por flexibilidade, inovação e valorização do fator humano, como evidenciado pela transição histórica para a Teoria das Relações Humanas e pelas críticas à rigidez burocrática.

Os avanços tecnológicos, especialmente em inteligência artificial, big data e automação, trazem à tona o conceito de “Taylorismo Digital”. Plataformas digitais monitoram cada clique, tempo de resposta e produtividade de trabalhadores, replicando a lógica de controle e eficiência defendida por Taylor, agora em escala ampliada e mediada por softwares. Empresas do setor de logística e de aplicativos de transporte aplicam métricas constantes de desempenho, estabelecendo padrões de comportamento e reduzindo margens de autonomia dos trabalhadores.

Assim, se por um lado a tecnologia possibilita ganhos de produtividade e precisão, por outro, recria dilemas clássicos relacionados à alienação, repetição e controle excessivo. O desafio contemporâneo consiste em equilibrar eficiência digital com autonomia, criatividade e bem-estar humano, para que o Taylorismo Digital não se transforme em uma “jaula de ferro” moderna, mas sim em um meio de potencializar o trabalho humano em um ambiente mais justo e inovador.
'''

for paragrafo in texto.split('\n\n'):
    p = doc.add_paragraph(paragrafo)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Numeração de páginas
section = doc.sections[-1]
footer = section.footer
p_footer = footer.paragraphs[0]
p_footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p_footer.add_run("Página ")
from docx.oxml import OxmlElement
fldChar = OxmlElement('w:fldChar')
fldChar.set('w:fldCharType', 'begin')
p_footer._p.append(fldChar)
fldCharSep = OxmlElement('w:fldChar')
fldCharSep.set('w:fldCharType', 'separate')
p_footer._p.append(fldCharSep)
fldCharEnd = OxmlElement('w:fldChar')
fldCharEnd.set('w:fldCharType', 'end')
p_footer._p.append(fldCharEnd)

# Salvar arquivo
doc.save('/mnt/data/Taylorismo_Digital_Senac.docx')
